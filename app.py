import streamlit as st
import time
import io
import re
from datetime import datetime


def generate_pdf(topic: str, report: str, feedback: str) -> bytes:
    from fpdf import FPDF

    def clean_text(text):
        # Convert markdown links [text](url) -> "text (url)"
        text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1 (\2)", text)
        text = text.replace("**", "").replace("__", "")
        return text.encode("latin-1", "replace").decode("latin-1")

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    eff_width = 190

    # ── Header banner ──
    pdf.set_fill_color(245, 240, 232)
    pdf.rect(0, 0, 210, 45, "F")
    pdf.set_font("Helvetica", "B", 20)
    pdf.set_text_color(90, 75, 60)
    pdf.set_xy(10, 10)
    pdf.multi_cell(eff_width, 10, "Research Report", align="C")
    pdf.set_font("Helvetica", "", 13)
    pdf.set_text_color(130, 110, 90)
    pdf.multi_cell(eff_width, 8, clean_text(topic), align="C")
    pdf.set_font("Helvetica", "", 9)
    pdf.set_text_color(160, 150, 140)
    pdf.cell(eff_width, 8, f"Generated: {datetime.now().strftime('%B %d, %Y  %H:%M')}", align="R", ln=True)
    pdf.set_y(50)
    pdf.set_font("Helvetica", "B", 13)
    pdf.set_text_color(80, 65, 50)
    pdf.cell(eff_width, 8, "Report", ln=True)
    pdf.set_draw_color(210, 195, 175)
    pdf.line(10, pdf.get_y(), 200, pdf.get_y())
    pdf.ln(5)
    pdf.set_font("Helvetica", "", 10)
    pdf.set_text_color(60, 55, 50)

    def write_body(pdf, text):
        """Render markdown-ish text with headers, bullets, and numbered lists."""
        for line in text.split("\n"):
            stripped = line.strip()
            if not stripped:
                pdf.ln(3)
                continue

            c_line = clean_text(stripped)

            if stripped.startswith("### "):
                pdf.ln(2)
                pdf.set_font("Helvetica", "B", 11)
                pdf.set_text_color(100, 85, 65)
                pdf.multi_cell(eff_width, 6.5, c_line.lstrip("# ").strip())
                pdf.set_font("Helvetica", "", 10)
                pdf.set_text_color(60, 55, 50)

            elif stripped.startswith("## "):
                pdf.ln(2)
                pdf.set_font("Helvetica", "B", 12)
                pdf.set_text_color(90, 75, 55)
                pdf.multi_cell(eff_width, 7, c_line.lstrip("# ").strip())
                pdf.set_font("Helvetica", "", 10)
                pdf.set_text_color(60, 55, 50)

            elif stripped.startswith("# "):
                pdf.ln(3)
                pdf.set_font("Helvetica", "B", 13)
                pdf.set_text_color(80, 65, 50)
                pdf.multi_cell(eff_width, 7, c_line.lstrip("# ").strip())
                pdf.set_font("Helvetica", "", 10)
                pdf.set_text_color(60, 55, 50)

            elif stripped.startswith("- ") or stripped.startswith("* "):
                bullet_text = c_line.lstrip("-*").strip()
                pdf.set_x(pdf.l_margin + 4)
                pdf.multi_cell(eff_width - 4, 6, f"-  {bullet_text}")

            elif re.match(r"^\d+\.\s", stripped):
                pdf.set_x(pdf.l_margin + 4)
                pdf.multi_cell(eff_width - 4, 6, c_line)

            else:
                pdf.multi_cell(eff_width, 6, c_line)

    write_body(pdf, report)

    # ── Critic Feedback section: only start a new page if there isn't enough room ──
    PAGE_HEIGHT = 297
    BOTTOM_MARGIN = 15
    MIN_SPACE_NEEDED = 40  # mm needed for a heading + a couple lines

    if pdf.get_y() > (PAGE_HEIGHT - BOTTOM_MARGIN - MIN_SPACE_NEEDED):
        pdf.add_page()
        pdf.set_fill_color(245, 240, 232)
        pdf.rect(0, 0, 210, 18, "F")
        pdf.set_font("Helvetica", "B", 13)
        pdf.set_text_color(80, 65, 50)
        pdf.set_xy(10, 5)
        pdf.cell(eff_width, 8, "Critic Feedback")
        pdf.ln(15)
    else:
        pdf.ln(8)
        pdf.set_draw_color(210, 195, 175)
        pdf.line(10, pdf.get_y(), 200, pdf.get_y())
        pdf.ln(6)
        pdf.set_font("Helvetica", "B", 13)
        pdf.set_text_color(80, 65, 50)
        pdf.cell(eff_width, 8, "Critic Feedback", ln=True)
        pdf.ln(4)

    pdf.set_font("Helvetica", "", 10)
    pdf.set_text_color(60, 55, 50)
    write_body(pdf, feedback)

    return bytes(pdf.output())


def generate_docx(topic: str, report: str, feedback: str) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    def clean(text):
        text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1 (\2)", text)
        return text.replace("**", "").replace("__", "")

    doc = Document()
    title = doc.add_heading(level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Research Report")
    run.font.color.rgb = RGBColor(90, 75, 60)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run(topic)
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(130, 110, 90)
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    mr = meta.add_run(f"Generated: {datetime.now().strftime('%B %d, %Y  %H:%M')}")
    mr.font.size = Pt(9)
    mr.font.color.rgb = RGBColor(160, 150, 140)
    doc.add_paragraph()

    def write_body(text):
        for line in text.split("\n"):
            stripped = line.strip()
            if not stripped:
                doc.add_paragraph()
                continue

            c_line = clean(stripped)

            if stripped.startswith("### "):
                doc.add_heading(c_line.lstrip("# ").strip(), level=3)

            elif stripped.startswith("## "):
                h = doc.add_heading(c_line.lstrip("# ").strip(), level=2)
                if h.runs:
                    h.runs[0].font.color.rgb = RGBColor(100, 85, 65)

            elif stripped.startswith("# "):
                h = doc.add_heading(c_line.lstrip("# ").strip(), level=1)
                if h.runs:
                    h.runs[0].font.color.rgb = RGBColor(80, 65, 50)

            elif stripped.startswith("- ") or stripped.startswith("* "):
                bullet_text = c_line.lstrip("-*").strip()
                doc.add_paragraph(bullet_text, style="List Bullet")

            elif re.match(r"^\d+\.\s", stripped):
                doc.add_paragraph(c_line, style="List Number")

            else:
                p = doc.add_paragraph(c_line)
                if p.runs:
                    p.runs[0].font.size = Pt(11)

    write_body(report)

    # No forced page break — Word will flow naturally.
    doc.add_paragraph()
    ch = doc.add_heading("Critic Feedback", level=1)
    if ch.runs:
        ch.runs[0].font.color.rgb = RGBColor(80, 65, 50)

    write_body(feedback)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def generate_txt(topic: str, report: str, feedback: str) -> bytes:
    divider = "=" * 60
    content = f"""RESEARCH REPORT
{divider}
Topic  : {topic}
Date   : {datetime.now().strftime('%B %d, %Y  %H:%M')}
{divider}

{report}

{divider}
CRITIC FEEDBACK
{divider}

{feedback}
"""
    return content.encode("utf-8")


st.set_page_config(
    page_title="Research Pipeline",
    page_icon="📜",
    layout="wide",
    initial_sidebar_state="collapsed",
)

st.markdown("""
<style>
  /* ── Google Font Import ── */
  @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700;800&display=swap');

  /* ══════════════════════════════════════
     BASE & BACKGROUND
  ══════════════════════════════════════ */
  [data-testid="stAppViewContainer"],
  [data-testid="stAppViewContainer"] > .main,
  .main,
  [data-testid="stHeader"],
  [data-testid="stToolbar"],
  [data-testid="stDecoration"],
  [data-testid="stBottom"],
  section[data-testid="stSidebar"] {
    background-color: #080b12 !important;
  }

  /* ══════════════════════════════════════
     GLOBAL FONT & TEXT
  ══════════════════════════════════════ */
  html, body, [class*="css"] {
    font-family: 'Inter', 'Georgia', sans-serif !important;
    color: #e2e8f0 !important;
  }

  /* ══════════════════════════════════════
     HEADINGS
  ══════════════════════════════════════ */
  h1 {
    color: #ffffff !important;
    font-weight: 800 !important;
    letter-spacing: -0.5px !important;
  }
  h2, h3, h4, h5, h6 {
    color: #f1f5f9 !important;
    font-weight: 700 !important;
  }

  /* ══════════════════════════════════════
     PARAGRAPHS & GENERAL TEXT
  ══════════════════════════════════════ */
  p, li, span, label {
    color: #cbd5e1 !important;
  }
  [data-testid="stMarkdownContainer"] p,
  [data-testid="stMarkdownContainer"] li,
  [data-testid="stMarkdownContainer"] span,
  [data-testid="stMarkdownContainer"] strong,
  [data-testid="stMarkdownContainer"] a {
    color: #cbd5e1 !important;
  }
  [data-testid="stMarkdownContainer"] strong {
    color: #f8fafc !important;
    font-weight: 700 !important;
  }

  /* ══════════════════════════════════════
     HERO ANIMATIONS
  ══════════════════════════════════════ */
  @keyframes gradientShift {
    0%   { background-position: 0% 50%; }
    50%  { background-position: 100% 50%; }
    100% { background-position: 0% 50%; }
  }
  @keyframes iconPulse {
    0%, 100% { transform: scale(1);   filter: drop-shadow(0 0 12px #d4a84388); }
    50%       { transform: scale(1.1); filter: drop-shadow(0 0 28px #d4a843cc); }
  }
  @keyframes borderGlow {
    0%, 100% { border-color: #1e293b; box-shadow: 0 0 0px #d4a84300; }
    50%       { border-color: #d4a84366; box-shadow: 0 0 24px #d4a84322; }
  }
  @keyframes badgeFadeIn {
    from { opacity: 0; transform: translateY(6px); }
    to   { opacity: 1; transform: translateY(0); }
  }
  @keyframes floatDot {
    0%, 100% { transform: translateY(0px);  opacity: 0.4; }
    50%       { transform: translateY(-8px); opacity: 1; }
  }

  /* ══════════════════════════════════════
     HERO WRAPPER
  ══════════════════════════════════════ */
  .hero-wrapper {
    text-align: center;
    padding: 48px 20px 16px;
    position: relative;
  }

  /* floating dots decoration */
  .hero-wrapper::before,
  .hero-wrapper::after {
    content: '●';
    position: absolute;
    font-size: 8px;
    color: #d4a843;
    animation: floatDot 3s ease-in-out infinite;
  }
  .hero-wrapper::before { top: 30px; left: 18%; animation-delay: 0s; }
  .hero-wrapper::after  { top: 50px; right: 18%; animation-delay: 1.5s; }

  /* glowing icon */
  .hero-icon {
    font-size: 52px;
    display: inline-block;
    animation: iconPulse 3s ease-in-out infinite;
    margin-bottom: 6px;
  }

  /* animated gradient title */
  .hero-title {
    font-size: 38px;
    font-weight: 800;
    font-family: 'Inter', sans-serif;
    letter-spacing: -1px;
    background: linear-gradient(270deg, #ffffff, #d4a843, #f0c050, #ffffff, #a0c4ff);
    background-size: 400% 400%;
    -webkit-background-clip: text;
    -webkit-text-fill-color: transparent;
    background-clip: text;
    animation: gradientShift 6s ease infinite;
    margin: 4px 0 6px;
    line-height: 1.15;
  }

  /* subtitle */
  .hero-sub {
    color: #64748b;
    font-size: 15px;
    font-family: 'Inter', sans-serif;
    font-weight: 400;
    max-width: 500px;
    margin: 0 auto 18px;
    line-height: 1.6;
  }

  /* badge row */
  .hero-badges {
    display: flex;
    justify-content: center;
    gap: 10px;
    flex-wrap: wrap;
    margin-bottom: 6px;
    animation: badgeFadeIn 0.8s ease both;
    animation-delay: 0.3s;
  }
  .hero-badge {
    background: #0d1117;
    border: 1px solid #1e293b;
    border-radius: 99px;
    padding: 4px 14px;
    font-size: 12px;
    font-weight: 600;
    color: #94a3b8;
    font-family: 'Inter', sans-serif;
    letter-spacing: 0.2px;
  }
  .hero-badge.gold {
    border-color: #d4a84366;
    color: #d4a843;
    background: #0f0d00;
  }

  /* divider with glow */
  .hero-divider {
    border: none;
    border-top: 1px solid #1e293b;
    margin: 24px 0;
    position: relative;
  }
  .hero-divider::after {
    content: '';
    position: absolute;
    top: -1px;
    left: 50%;
    transform: translateX(-50%);
    width: 120px;
    height: 1px;
    background: linear-gradient(90deg, transparent, #d4a843, transparent);
  }

  /* ══════════════════════════════════════
     TEXT INPUT
  ══════════════════════════════════════ */
  .stTextInput > div > div > input {
    background: #0f1420 !important;
    border: 1.5px solid #1e293b !important;
    border-radius: 10px !important;
    color: #f8fafc !important;
    font-size: 15px !important;
    padding: 11px 16px !important;
    font-family: 'Inter', sans-serif !important;
    transition: border-color 0.2s, box-shadow 0.2s !important;
  }
  .stTextInput > div > div > input::placeholder {
    color: #475569 !important;
  }
  .stTextInput > div > div > input:focus {
    border-color: #d4a843 !important;
    box-shadow: 0 0 0 3px #d4a84325 !important;
    outline: none !important;
  }

  /* ══════════════════════════════════════
     RUN PIPELINE BUTTON
  ══════════════════════════════════════ */
  div[data-testid="stButton"] > button {
    background: linear-gradient(135deg, #d4a843 0%, #b8860b 100%) !important;
    color: #000000 !important;
    border: none !important;
    border-radius: 10px !important;
    font-weight: 800 !important;
    font-size: 15px !important;
    padding: 11px 28px !important;
    letter-spacing: 0.3px !important;
    box-shadow: 0 0 18px #d4a84355, 0 2px 8px #00000055 !important;
    transition: all 0.2s ease !important;
    font-family: 'Inter', sans-serif !important;
  }
  div[data-testid="stButton"] > button:hover {
    background: linear-gradient(135deg, #f0c050 0%, #d4a843 100%) !important;
    box-shadow: 0 0 28px #d4a84388, 0 4px 16px #00000066 !important;
    transform: translateY(-2px) !important;
    color: #000000 !important;
  }
  div[data-testid="stButton"] > button:active {
    transform: translateY(0px) !important;
  }
  div[data-testid="stButton"] > button:disabled {
    background: #1e293b !important;
    color: #475569 !important;
    box-shadow: none !important;
    cursor: not-allowed !important;
  }

  /* ══════════════════════════════════════
     DOWNLOAD BUTTONS
  ══════════════════════════════════════ */
  div[data-testid="stDownloadButton"] > button {
    background: #0f1420 !important;
    color: #d4a843 !important;
    border: 1.5px solid #1e3a5f !important;
    border-radius: 8px !important;
    font-size: 13px !important;
    font-weight: 700 !important;
    padding: 8px 18px !important;
    transition: all 0.2s ease !important;
  }
  div[data-testid="stDownloadButton"] > button:hover {
    background: #162032 !important;
    border-color: #d4a843 !important;
    box-shadow: 0 0 14px #d4a84333 !important;
    color: #f0c050 !important;
  }

  /* ══════════════════════════════════════
     PROGRESS BAR
  ══════════════════════════════════════ */
  .stProgress > div > div > div {
    background: #0f1420 !important;
    border-radius: 99px !important;
  }
  .stProgress > div > div > div > div {
    background: linear-gradient(90deg, #d4a843, #f0c050, #d4a843) !important;
    background-size: 200% !important;
    animation: shimmer 2s infinite linear !important;
    border-radius: 99px !important;
  }
  @keyframes shimmer {
    0%   { background-position: 0% center; }
    100% { background-position: 200% center; }
  }
  .stProgress p {
    color: #d4a843 !important;
    font-weight: 600 !important;
    font-size: 13px !important;
  }

  /* ══════════════════════════════════════
     TABS
  ══════════════════════════════════════ */
  [data-baseweb="tab-list"] {
    background: #080b12 !important;
    border-bottom: 1.5px solid #1e293b !important;
    gap: 4px !important;
  }
  button[data-baseweb="tab"] {
    background: transparent !important;
    font-size: 15px !important;
    font-weight: 600 !important;
    color: #64748b !important;
    padding: 10px 20px !important;
    border-radius: 8px 8px 0 0 !important;
    transition: all 0.2s ease !important;
    font-family: 'Inter', sans-serif !important;
    letter-spacing: 0.2px !important;
  }
  button[data-baseweb="tab"]:hover {
    color: #cbd5e1 !important;
    background: #0f1420 !important;
  }
  button[data-baseweb="tab"][aria-selected="true"] {
    color: #d4a843 !important;
    border-bottom: 2.5px solid #d4a843 !important;
    background: transparent !important;
    font-weight: 700 !important;
  }
  [data-baseweb="tab-panel"] {
    background: #080b12 !important;
    padding-top: 20px !important;
  }

  /* ══════════════════════════════════════
     TAB SECTION HEADINGS
  ══════════════════════════════════════ */
  .tab-section-heading {
    font-size: 26px !important;
    font-weight: 800 !important;
    color: #ffffff !important;
    letter-spacing: -0.3px !important;
    margin-bottom: 18px !important;
    padding-bottom: 10px !important;
    border-bottom: 2px solid #1e293b !important;
    font-family: 'Inter', sans-serif !important;
    display: flex !important;
    align-items: center !important;
    gap: 10px !important;
  }
  .tab-section-heading .accent-line {
    display: inline-block;
    width: 5px;
    height: 28px;
    background: linear-gradient(180deg, #d4a843, #b8860b);
    border-radius: 3px;
    margin-right: 4px;
    vertical-align: middle;
  }

  /* ══════════════════════════════════════
     METRICS
  ══════════════════════════════════════ */
  [data-testid="stMetric"] {
    background: #0f1420 !important;
    border: 1px solid #1e293b !important;
    border-radius: 12px !important;
    padding: 14px 18px !important;
  }
  [data-testid="stMetricLabel"] p {
    color: #64748b !important;
    font-size: 12px !important;
    font-weight: 600 !important;
    text-transform: uppercase !important;
    letter-spacing: 0.5px !important;
  }
  [data-testid="stMetricValue"] {
    color: #f8fafc !important;
    font-size: 24px !important;
    font-weight: 800 !important;
  }

  /* ══════════════════════════════════════
     ALERT / INFO / WARNING / ERROR
  ══════════════════════════════════════ */
  div[data-testid="stAlert"] {
    background: #0f1420 !important;
    border-radius: 10px !important;
    border: 1px solid #1e293b !important;
  }
  div[data-testid="stAlert"] p,
  div[data-testid="stAlert"] span {
    color: #cbd5e1 !important;
  }
  div[data-testid="stAlert"][kind="info"],
  div[class*="stInfo"] {
    border-left: 4px solid #3b82f6 !important;
    background: #0a1628 !important;
  }
  div[data-testid="stAlert"][kind="warning"],
  div[class*="stWarning"] {
    border-left: 4px solid #d4a843 !important;
    background: #12100a !important;
  }
  div[data-testid="stAlert"][kind="error"],
  div[class*="stError"] {
    border-left: 4px solid #ef4444 !important;
    background: #12080a !important;
  }

  /* ══════════════════════════════════════
     CODE BLOCK
  ══════════════════════════════════════ */
  pre, code, .stCode, [data-testid="stCode"] {
    background: #0a0d16 !important;
    color: #94a3b8 !important;
    border: 1px solid #1e293b !important;
    border-radius: 10px !important;
    font-size: 13px !important;
  }

  /* ══════════════════════════════════════
     DIVIDER
  ══════════════════════════════════════ */
  hr {
    border: none !important;
    border-top: 1px solid #1e293b !important;
  }

  /* ══════════════════════════════════════
     SCROLLBAR
  ══════════════════════════════════════ */
  ::-webkit-scrollbar { width: 5px; height: 5px; }
  ::-webkit-scrollbar-track { background: #080b12; }
  ::-webkit-scrollbar-thumb { background: #1e293b; border-radius: 99px; }
  ::-webkit-scrollbar-thumb:hover { background: #d4a843; }

  /* ══════════════════════════════════════
     HIDE STREAMLIT CHROME
  ══════════════════════════════════════ */
  #MainMenu, footer, header { visibility: hidden; }
</style>
""", unsafe_allow_html=True)

AGENTS = [
    {"id": "search", "name": "Search Agent", "icon": "🔍", "desc": "Finds recent, reliable sources on the topic"},
    {"id": "reader", "name": "Reader Agent", "icon": "📖", "desc": "Scrapes the most relevant URL for deeper content"},
    {"id": "writer", "name": "Writer Agent", "icon": "✍️",  "desc": "Drafts a comprehensive, structured report"},
    {"id": "critic", "name": "Critic Agent", "icon": "🧐", "desc": "Reviews the report for quality and accuracy"},
]

DEFAULTS = {
    "running": False,
    "current_step": -1,
    "results": {},
    "logs": [],
    "done": False,
    "error": None,
    "topic_locked": "",
}
for k, v in DEFAULTS.items():
    if k not in st.session_state:
        st.session_state[k] = v


def reset():
    for k, v in DEFAULTS.items():
        st.session_state[k] = v


def add_log(msg: str):
    ts = datetime.now().strftime("%H:%M:%S")
    st.session_state.logs.append(f"[{ts}]  {msg}")


def agent_card_html(agent, card_state):
    styles = {
        "idle":    ("#1e293b", "#94a3b8", "Waiting",    "#0f1420", "#64748b",  "1"),
        "active":  ("#d4a843", "#fde68a", "Running…",   "#1a1400", "#d4a843",  "1"),
        "done":    ("#22c55e", "#86efac", "Complete ✓", "#0a1f0a", "#4ade80",  "1"),
        "waiting": ("#111827", "#374151", "Waiting",    "#080b12", "#374151",  "0.45"),
    }
    border, name_color, status_text, badge_bg, badge_color, opacity = styles[card_state]
    shadow = "0 0 20px #d4a84322, 0 2px 8px #00000044" if card_state == "active" else "0 1px 4px #00000033"
    card_bg = "#0d1117" if card_state != "waiting" else "#090c13"
    return f"""
    <div style="background:{card_bg};border:1.5px solid {border};border-radius:12px;
      padding:15px 18px;margin-bottom:9px;opacity:{opacity};
      box-shadow:{shadow};transition:all .3s ease;">
      <div style="display:flex;justify-content:space-between;align-items:center">
        <span style="font-size:16px;color:{name_color};font-weight:700;
          font-family:'Inter',sans-serif">{agent['icon']} {agent['name']}</span>
        <span style="background:{badge_bg};color:{badge_color};border:1px solid {border};
          padding:3px 12px;border-radius:99px;font-size:11px;font-weight:700;
          font-family:'Inter',sans-serif;letter-spacing:0.3px">{status_text}</span>
      </div>
      <div style="color:#475569;font-size:13px;margin-top:5px;
        font-family:'Inter',sans-serif">{agent['desc']}</div>
    </div>
    """


# ── Creative Hero Header ─────────────────────────────────────────────────────
st.markdown("""
<div class="hero-wrapper">

  <!-- floating accent dots handled by CSS ::before / ::after -->

  <!-- glowing animated icon -->
  <div class="hero-icon">📜</div>

  <!-- animated gradient title -->
  <div class="hero-title">Research Pipeline</div>

  <!-- subtitle -->
  <div class="hero-sub">
    Four specialised AI agents — searching, reading, writing, and critiquing —
    working in sequence to deliver a polished research report in minutes.
  </div>

  <!-- feature badges -->
  <div class="hero-badges">
    <span class="hero-badge gold">⚡ AI-Powered</span>
    <span class="hero-badge">🔍 Live Search</span>
    <span class="hero-badge">✍️ Auto-Written</span>
    <span class="hero-badge">🧐 Critic-Reviewed</span>
    <span class="hero-badge">⬇ PDF · DOCX · TXT</span>
  </div>

</div>

<!-- glowing centre divider -->
<hr class="hero-divider">
""", unsafe_allow_html=True)

left, right = st.columns([1, 1.7], gap="large")

with left:
    topic_input = st.text_input(
        "Topic",
        placeholder="e.g. CRISPR gene editing, quantum computing, climate policy…",
        disabled=st.session_state.running,
        label_visibility="collapsed",
    )

    col_run, col_reset = st.columns([4, 1])
    with col_run:
        run_btn = st.button(
            "▶  Run Pipeline" if not st.session_state.running else "⏳  Pipeline running…",
            disabled=st.session_state.running or not topic_input.strip(),
            use_container_width=True,
        )
    with col_reset:
        if st.button("↺", help="Reset everything", use_container_width=True):
            reset()
            st.rerun()

    st.markdown("<br>", unsafe_allow_html=True)
    st.markdown("""
    <div style="font-size:13px;font-weight:700;color:#64748b;text-transform:uppercase;
      letter-spacing:1px;margin-bottom:10px;font-family:'Inter',sans-serif">
      🤖 &nbsp;Agent Status
    </div>
    """, unsafe_allow_html=True)

    step = st.session_state.current_step
    idle = not st.session_state.running and not st.session_state.done

    for i, agent in enumerate(AGENTS):
        if idle:
            state = "idle"
        elif st.session_state.done or i < step:
            state = "done"
        elif i == step:
            state = "active"
        else:
            state = "waiting"
        st.markdown(agent_card_html(agent, state), unsafe_allow_html=True)

    if st.session_state.running or st.session_state.done:
        total = len(AGENTS)
        done_count = step if st.session_state.running else total
        label = "✅ Complete!" if st.session_state.done else f"Step {done_count + 1} of {total}"
        st.progress(min(done_count / total, 1.0), text=label)

    if st.session_state.logs:
        st.markdown("<br>", unsafe_allow_html=True)
        st.markdown("""
        <div style="font-size:13px;font-weight:700;color:#64748b;text-transform:uppercase;
          letter-spacing:1px;margin-bottom:8px;font-family:'Inter',sans-serif">
          🗒️ &nbsp;Activity Log
        </div>
        """, unsafe_allow_html=True)
        st.code("\n".join(st.session_state.logs[-25:]), language=None)

with right:
    if not st.session_state.done and not st.session_state.running:
        st.markdown("""
        <div style="text-align:center;padding:90px 30px;background:#0d1117;
          border:1.5px dashed #1e293b;border-radius:16px;margin-top:10px;">
          <div style="font-size:52px">🔬</div>
          <div style="font-size:20px;color:#f1f5f9;margin-top:16px;font-weight:700;
            font-family:'Inter',sans-serif">
            Results will appear here
          </div>
          <div style="font-size:14px;color:#475569;margin-top:8px;font-family:'Inter',sans-serif">
            Enter a research topic on the left and press Run Pipeline
          </div>
        </div>
        """, unsafe_allow_html=True)

    elif st.session_state.error:
        st.error(f"Pipeline error: {st.session_state.error}")

    else:
        r = st.session_state.results
        locked_topic = st.session_state.topic_locked

        tab_s, tab_r, tab_w, tab_c, tab_dl = st.tabs([
            "🔍 Search", "📖 Reader", "✍️ Report", "🧐 Critic", "⬇ Download"
        ])

        with tab_s:
            if "search" in r:
                st.markdown("""
                <div class="tab-section-heading">
                  <span class="accent-line"></span>Search Results
                </div>""", unsafe_allow_html=True)
                st.markdown(
                    f"<div style='background:#0d1117;border:1px solid #1e293b;border-radius:12px;"
                    f"padding:20px;font-size:14px;color:#cbd5e1;white-space:pre-wrap;"
                    f"font-family:\"Inter\",sans-serif;line-height:1.7'>{r['search']}</div>",
                    unsafe_allow_html=True,
                )
            else:
                st.info("Search agent hasn't run yet.")

        with tab_r:
            if "reader" in r:
                st.markdown("""
                <div class="tab-section-heading">
                  <span class="accent-line"></span>Scraped Content
                </div>""", unsafe_allow_html=True)
                st.markdown(
                    f"<div style='background:#0d1117;border:1px solid #1e293b;border-radius:12px;"
                    f"padding:20px;font-size:14px;color:#cbd5e1;white-space:pre-wrap;"
                    f"font-family:\"Inter\",sans-serif;line-height:1.7'>{r['reader']}</div>",
                    unsafe_allow_html=True,
                )
            else:
                st.info("Reader agent hasn't run yet.")

        with tab_w:
            if "report" in r:
                st.markdown("""
                <div class="tab-section-heading">
                  <span class="accent-line"></span>Final Report
                </div>""", unsafe_allow_html=True)
                c1, c2 = st.columns(2)
                c1.metric("Words", f"{len(r['report'].split()):,}")
                c2.metric("Topic", locked_topic[:28] + "…" if len(locked_topic) > 28 else locked_topic)
                st.markdown("<br>", unsafe_allow_html=True)
                st.markdown(r["report"])
            else:
                st.info("Writer agent hasn't run yet.")

        with tab_c:
            if "feedback" in r:
                st.markdown("""
                <div class="tab-section-heading">
                  <span class="accent-line"></span>Critic Feedback
                </div>""", unsafe_allow_html=True)
                st.markdown(
                    f"<div style='background:#0d1117;border:1px solid #1e293b;border-radius:12px;"
                    f"padding:20px;font-size:14px;color:#cbd5e1;white-space:pre-wrap;"
                    f"font-family:\"Inter\",sans-serif;line-height:1.7'>{r['feedback']}</div>",
                    unsafe_allow_html=True,
                )
            else:
                st.info("Critic agent hasn't run yet.")

        with tab_dl:
            if "report" in r and "feedback" in r:
                st.markdown("""
                <div class="tab-section-heading">
                  <span class="accent-line"></span>Download Report
                </div>""", unsafe_allow_html=True)
                st.markdown(
                    "<p style='color:#64748b;font-size:14px;font-family:\"Inter\",sans-serif;"
                    "margin-top:-8px;margin-bottom:20px'>"
                    "Choose a format to download the full report including critic feedback.</p>",
                    unsafe_allow_html=True,
                )

                fname_base = locked_topic.replace(" ", "_")[:40]
                col_pdf, col_docx, col_txt = st.columns(3)

                with col_pdf:
                    st.markdown(
                        "<div style='font-weight:700;color:#f1f5f9;font-family:\"Inter\",sans-serif;"
                        "font-size:15px;margin-bottom:4px'>📄 PDF</div>"
                        "<div style='font-size:12px;color:#475569;font-family:\"Inter\",sans-serif;"
                        "margin-bottom:10px'>Styled, print-ready</div>",
                        unsafe_allow_html=True
                    )
                    try:
                        pdf_bytes = generate_pdf(locked_topic, r["report"], r["feedback"])
                        st.download_button("Download PDF", data=pdf_bytes,
                            file_name=f"{fname_base}_report.pdf",
                            mime="application/pdf", use_container_width=True)
                    except Exception as e:
                        st.warning(f"Error generating PDF: {e}")

                with col_docx:
                    st.markdown(
                        "<div style='font-weight:700;color:#f1f5f9;font-family:\"Inter\",sans-serif;"
                        "font-size:15px;margin-bottom:4px'>📝 DOCX</div>"
                        "<div style='font-size:12px;color:#475569;font-family:\"Inter\",sans-serif;"
                        "margin-bottom:10px'>Editable Word doc</div>",
                        unsafe_allow_html=True
                    )
                    try:
                        docx_bytes = generate_docx(locked_topic, r["report"], r["feedback"])
                        st.download_button("Download DOCX", data=docx_bytes,
                            file_name=f"{fname_base}_report.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True)
                    except Exception as e:
                        st.warning(f"Error generating DOCX: {e}")

                with col_txt:
                    st.markdown(
                        "<div style='font-weight:700;color:#f1f5f9;font-family:\"Inter\",sans-serif;"
                        "font-size:15px;margin-bottom:4px'>🗒️ TXT</div>"
                        "<div style='font-size:12px;color:#475569;font-family:\"Inter\",sans-serif;"
                        "margin-bottom:10px'>Plain text, always works</div>",
                        unsafe_allow_html=True
                    )
                    txt_bytes = generate_txt(locked_topic, r["report"], r["feedback"])
                    st.download_button("Download TXT", data=txt_bytes,
                        file_name=f"{fname_base}_report.txt",
                        mime="text/plain", use_container_width=True)

                st.markdown("<br>", unsafe_allow_html=True)
                st.info("Note: PDF/DOCX require `fpdf2` and `python-docx`. (Ignore if already installed)")
            else:
                st.markdown("""
                <div style="text-align:center;padding:50px 20px;background:#0d1117;
                  border:1.5px dashed #1e293b;border-radius:12px">
                  <div style="font-size:36px">⏳</div>
                  <div style="color:#475569;margin-top:12px;font-family:'Inter',sans-serif;
                    font-size:15px;font-weight:500">
                    Downloads available once the pipeline completes.
                  </div>
                </div>
                """, unsafe_allow_html=True)


# ── Pipeline wiring ──────────────────────────────────────────────────────────

# ── Pipeline wiring ──────────────────────────────────────────────────────────

# ── Pipeline wiring ──────────────────────────────────────────

def _run_step(step_idx: int, topic: str, results: dict):
    from tools import web_search, scrape_url
    from pipeline import writer_chain, critic_chain
    import re

    # ✅ STEP 1 — SEARCH
    if step_idx == 0:
        return web_search.invoke({"query": topic})

    # ✅ STEP 2 — READER (Extract first URL & scrape)
    elif step_idx == 1:
        search_text = results.get("search", "")
        urls = re.findall(r'https?://\S+', search_text)

        if urls:
            best_url = urls[0]
            return scrape_url.invoke({"url": best_url})
        else:
            return "No valid URL found in search results."

    # ✅ STEP 3 — WRITER (Gemini)
    elif step_idx == 2:
        combined = (
            f"SEARCH RESULTS:\n{results.get('search', '')}\n\n"
            f"DETAILED SCRAPED CONTENT:\n{results.get('reader', '')}"
        )

        return writer_chain.invoke({
            "topic": topic,
            "research": combined
        })

    # ✅ STEP 4 — CRITIC (Gemini)
    elif step_idx == 3:
        return critic_chain.invoke({
            "report": results.get("report", "")
        })


STEP_KEYS  = ["search", "reader", "report", "feedback"]
STEP_NAMES = ["Search Agent", "Reader Agent", "Writer Agent", "Critic Agent"]


# ✅ START PIPELINE WHEN BUTTON CLICKED
if run_btn and topic_input.strip():
    reset()
    st.session_state.running = True
    st.session_state.current_step = 0
    st.session_state.topic_locked = topic_input.strip()
    add_log(f"Pipeline started → '{topic_input.strip()}'")
    st.rerun()


# ✅ EXECUTE PIPELINE STEP-BY-STEP
if st.session_state.running:
    i = st.session_state.current_step
    topic_locked = st.session_state.topic_locked

    if 0 <= i < len(STEP_KEYS):
        key  = STEP_KEYS[i]
        name = STEP_NAMES[i]

        if key not in st.session_state.results:
            try:
                add_log(f"{name} started…")

                result = _run_step(i, topic_locked, st.session_state.results)

                st.session_state.results[key] = result
                add_log(f"{name} done ✓")

                if i + 1 < len(STEP_KEYS):
                    st.session_state.current_step = i + 1
                else:
                    st.session_state.running = False
                    st.session_state.done = True
                    st.session_state.current_step = len(STEP_KEYS)
                    add_log("All agents complete — report ready!")

                st.rerun()

            except Exception as e:
                st.session_state.error = str(e)
                st.session_state.running = False
                add_log(f"Error in {name}: {e}")
                st.rerun()
